from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
ASSETS = ROOT / ".tmp" / "slr_pattern" / "generated_figures"
OUT.mkdir(exist_ok=True)
ASSETS.mkdir(parents=True, exist_ok=True)

BLUE = "1F4E79"
TEXT = "1F1F1F"
GRAY = "555555"
PALE = "D9EAF7"
PALE_2 = "EAF2F8"
GREEN = "D9EAD3"
GOLD = "FCE5CD"
WHITE = "FFFFFF"


def pil_colour(value: str):
    return value if value.startswith("#") else f"#{value}"


def arial(size: int, bold: bool = False):
    candidate = "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"
    return ImageFont.truetype(candidate, size)


def centered(draw, box, text, font, fill=TEXT):
    left, top, right, bottom = box
    bounds = draw.multiline_textbbox((0, 0), text, font=font, spacing=4, align="center")
    x = left + (right - left - (bounds[2] - bounds[0])) / 2
    y = top + (bottom - top - (bounds[3] - bounds[1])) / 2
    draw.multiline_text((x, y), text, font=font, fill=pil_colour(fill), spacing=4, align="center")


def box(draw, coords, text, fill, *, title=False):
    draw.rounded_rectangle(coords, radius=22, fill=pil_colour(fill), outline=pil_colour(BLUE), width=3)
    centered(draw, coords, text, arial(27 if title else 24, bold=title))


def arrow(draw, start, end):
    draw.line([start, end], fill=pil_colour(BLUE), width=5)
    x1, y1 = end
    draw.polygon([(x1, y1), (x1 - 18, y1 - 10), (x1 - 18, y1 + 10)], fill=pil_colour(BLUE))


def make_keyword_map(path: Path, left: str, centre: str, right: str):
    image = Image.new("RGB", (1600, 780), pil_colour(WHITE))
    draw = ImageDraw.Draw(image)
    box(draw, (540, 285, 1060, 495), centre, PALE, title=True)
    nodes = [
        ((80, 100, 460, 250), left, GREEN),
        ((1140, 100, 1520, 250), right, GOLD),
        ((80, 535, 460, 685), "Governance and\nlegal context", PALE_2),
        ((1140, 535, 1520, 685), "Evaluation and\nadoption", PALE_2),
    ]
    for rect, text, color in nodes:
        box(draw, rect, text, color)
        cx = (rect[0] + rect[2]) // 2
        cy = (rect[1] + rect[3]) // 2
        tx = 800 if cx < 800 else 800
        ty = 390 if cy < 390 else 390
        arrow(draw, (cx + (18 if cx < 800 else -18), cy), (tx + (-270 if cx < 800 else 270), ty))
    image.save(path)


def make_workflow(path: Path):
    image = Image.new("RGB", (1600, 430), pil_colour(WHITE))
    draw = ImageDraw.Draw(image)
    labels = ["Identify\nrecords", "Screen against\ncriteria", "Quality\nappraise", "Extract\nevidence", "Thematic\nsynthesis", "Report\nfindings"]
    colors = [PALE, GREEN, PALE_2, GOLD, PALE, GREEN]
    width, gap, top, height = 220, 35, 115, 180
    for idx, (label, color) in enumerate(zip(labels, colors)):
        left = 25 + idx * (width + gap)
        box(draw, (left, top, left + width, top + height), label, color)
        if idx < len(labels) - 1:
            arrow(draw, (left + width + 4, top + height // 2), (left + width + gap - 5, top + height // 2))
    centered(draw, (100, 15, 1500, 85), "PRISMA-aligned review workflow (reported without fabricated database counts)", arial(28, bold=True), BLUE)
    image.save(path)


def make_evidence_map(path: Path, rows: list[tuple[str, str, str]]):
    image = Image.new("RGB", (1600, 760), pil_colour(WHITE))
    draw = ImageDraw.Draw(image)
    title_font = arial(27, bold=True)
    body_font = arial(20)
    headers = [(50, 30, 500, 100, "Evidence theme"), (535, 30, 1045, 100, "What the literature indicates"), (1080, 30, 1550, 100, "BLMS implication")]
    for x1, y1, x2, y2, title in headers:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=12, fill=pil_colour(BLUE))
        centered(draw, (x1, y1, x2, y2), title, title_font, pil_colour(WHITE))
    top = 120
    for i, (a, b, c) in enumerate(rows):
        h = 145
        fill = PALE_2 if i % 2 == 0 else WHITE
        cols = [(50, 500, a), (535, 1045, b), (1080, 1550, c)]
        for x1, x2, text in cols:
            draw.rounded_rectangle((x1, top, x2, top + h), radius=10, fill=pil_colour(fill), outline="#A7B9C9", width=2)
            centered(draw, (x1 + 12, top + 8, x2 - 12, top + h - 8), text, body_font)
        top += h + 12
    image.save(path)


def set_run(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths: list[int], indent: int = 0):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            mar = tc_pr.first_child_found_in("w:tcMar")
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for side, value in (("top", "80"), ("bottom", "80"), ("start", "100"), ("end", "100")):
                node = mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    mar.append(node)
                node.set(qn("w:w"), value)
                node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    item = OxmlElement("w:tblHeader")
    item.set(qn("w:val"), "true")
    tr_pr.append(item)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run(run, 9, color=GRAY)


def doc_base(side_margin=1.0):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(side_margin)
    sec.right_margin = Inches(side_margin)
    sec.header_distance = Inches(0.49)
    sec.footer_distance = Inches(0.49)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(8)
    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, TEXT),
        ("Heading 2", 16, 18, 6, TEXT),
        ("Heading 3", 14, 16, 4, "434343"),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Page ")
    set_run(r, 9, color=GRAY)
    add_page_number(footer)
    return doc


def text(doc, value, *, style=None, bold_lead=None, italic=False):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if bold_lead and value.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run(r, bold=True)
        r = p.add_run(value[len(bold_lead):])
        set_run(r, italic=italic)
    else:
        r = p.add_run(value)
        set_run(r, italic=italic)
    return p


def heading(doc, value, level=3):
    return doc.add_heading(value, level=level)


def list_item(doc, value):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(value)
    set_run(r)
    return p


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_geometry(t, widths)
    repeat_header(t.rows[0])
    for cell, value in zip(t.rows[0].cells, headers):
        shade(cell, PALE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run(r, 9.2, bold=True, color=BLUE)
    for row in rows:
        cells = t.add_row().cells
        for cell, value in zip(cells, row):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(value)
            set_run(r, 9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def figure(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    shape = p.add_run().add_picture(str(image_path), width=Inches(6.3))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(caption)
    set_run(r, 9.5, italic=True, color=GRAY)


def add_method(doc, rqs, keywords, boolean, period):
    heading(doc, "2. PROCESS OF SLR", 2)
    heading(doc, "Planning the Review", 3)
    text(doc, "This review was planned to identify, evaluate, and synthesise scholarly evidence relevant to the BLMS research problem. The protocol was defined before drafting the findings: research questions, search strings, eligibility criteria, quality checks, extraction fields, and synthesis categories were specified to reduce ad-hoc selection, following the structured-review discipline of Kitchenham and Charters (2007) and the reporting principles of PRISMA 2020 (Page et al., 2021).")
    heading(doc, "Conducting the Review", 3)
    text(doc, "The conduct phase follows six connected steps used in software-engineering and information-systems reviews.")
    for label, content in [
        ("i. Research Questions (RQs): ", "The RQs translate the BLMS objectives into answerable evidence needs."),
        ("ii. Search Strategy (SS): ", "Search strings combine the land-administration, blockchain, identity, document-integrity, evaluation, and Nigeria context terms."),
        ("iii. Study Selection Criteria (SSC): ", "Titles, abstracts, and full text are screened against published eligibility rules."),
        ("iv. Quality Assessment Criteria (QAC): ", "Each retained source is checked for traceable publication details, clear context, relevant method or architecture, and stated limitations."),
        ("v. Data Extraction and Monitoring (DEM): ", "A structured sheet records authorship, setting, technology, workflow, evaluation approach, findings, and limitations."),
        ("vi. Data Synthesis (DS): ", "Findings are grouped by RQ to distinguish evidence about the problem from evidence about the proposed solution."),
    ]:
        text(doc, label + content, bold_lead=label)
    heading(doc, "Reporting the Review", 3)
    text(doc, "The review reports the protocol, included evidence themes, and limitations in a form that can be checked and extended by the supervisor. The report does not claim a live benchmark, a field deployment, or database-search totals that were not independently recorded.")

    text(doc, "TABLE 1. Define research questions for the SLR.", bold_lead="TABLE 1.")
    table(doc, ["QID", "Research Questions", "Objectives"], rqs, [720, 4200, 4440])
    heading(doc, "Search Strategy (SS)", 3)
    text(doc, "The search protocol covers peer-reviewed journal and conference literature, recognised technical reports, and institutional or publisher records that can be verified through a DOI, publisher page, or repository. The search period is " + period + ". Relevant records may be expanded through backward reference checking, provided they meet the same inclusion criteria.")
    for source in ["ScienceDirect (Elsevier)", "ACM Digital Library", "IEEE Xplore", "SpringerLink", "Wiley Online Library", "Scopus and Web of Science (where access is available)", "Google Scholar and institutional repositories for discovery and citation chaining"]:
        list_item(doc, source)
    heading(doc, "Search Keywords and Terms", 3)
    text(doc, "The core search concepts are: " + keywords + ".")
    heading(doc, "Boolean Search", 3)
    text(doc, "The Boolean search was designed to preserve the logical links between the project context, the technical intervention, and the evaluation/adoption context.")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    r = p.add_run(boolean)
    set_run(r, 9.5, color=BLUE)


def add_selection_and_quality(doc, flow, evidence, inclusion, exclusion):
    heading(doc, "Study Selection Process", 3)
    text(doc, "The study-selection process follows a PRISMA-aligned sequence: identify records, remove duplicates where applicable, screen titles and abstracts, assess full text against eligibility criteria, appraise quality, and synthesise the retained evidence. The workflow is shown without numerical claims because exact database export counts were not captured for this submission.")
    figure(doc, flow, "Fig. 2. PRISMA-aligned review workflow used for this BLMS SLR.")
    heading(doc, "Inclusion and Exclusion Criteria", 3)
    text(doc, "Inclusion criteria:", bold_lead="Inclusion criteria:")
    for item in inclusion:
        list_item(doc, item)
    text(doc, "Exclusion criteria:", bold_lead="Exclusion criteria:")
    for item in exclusion:
        list_item(doc, item)
    heading(doc, "Quality Evaluation", 3)
    text(doc, "The quality evaluation asks whether a source states a clear context and objective, provides a traceable publication record, describes its technical or empirical method, presents findings relevant to at least one RQ, and acknowledges limitations or constraints. Sources that are only promotional, lack identifiable authorship, or cannot be verified are excluded.")
    figure(doc, evidence, "Fig. 3. Evidence-synthesis map linking reviewed themes to BLMS design implications.")


def add_references(doc, entries):
    heading(doc, "REFERENCES", 3)
    for entry in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(entry)
        set_run(r, 9.5)


def build_problem():
    doc = doc_base(1.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("A Systematic Literature Review on the Problem Domain of Blockchain-Based Land Management in Nigeria")
    set_run(r, 18, bold=True)
    text(doc, "ABSTRACT", bold_lead="ABSTRACT")
    text(doc, "Land registration is expected to provide reliable evidence of rights, parcels, transactions, and official decisions. In Nigeria, the literature reports delays, weak records, high verification effort, opaque procedures, fragmented administrative practice, and risks of duplicate allocation or sale. This systematic literature review defines the problem domain for the Blockchain-Based Land Management System (BLMS). It synthesises research on Nigerian land administration, digital land records, verification, public-sector governance, and comparable low-resource contexts. The review is structured around five research questions covering the sources of ownership uncertainty, available records and datasets, evaluation measures, comparisons between traditional and digital systems, and implementation barriers. The synthesis finds that a technical platform can improve traceability and document-integrity checks, but it cannot independently cure inaccurate source records, legal ambiguity, institutional coordination failures, identity risk, or limited digital access. The resulting problem specification supports a phased BLMS that keeps land-authority approval, correction, privacy, and dispute-resolution requirements explicit.")
    text(doc, "Keywords: land administration; land registration; land title; Nigeria; land records; transparency; fraud; digital government.", bold_lead="Keywords:")
    heading(doc, "1. INTRODUCTION", 1)
    text(doc, "Land is both an economic asset and a source of social security. A credible land administration system must link a defined parcel, the persons or organisations claiming rights, supporting documents, and the public authority that recognises or records the relevant action. When these links are difficult to verify, citizens bear the cost through repeated searches, delayed registration, insecurity of tenure, and avoidable disputes. Nigerian studies identify weaknesses in registration procedures, documentary evidence, institutional capacity, and transparency as persistent concerns (Nwuba & Nuhu, 2018; Ibrahim et al., 2021; Okoli et al., 2024).")
    text(doc, "The problem domain is broader than a paper-versus-digital comparison. It includes source-data quality, identity, access control, document integrity, organisational accountability, legal authority, rural and urban connectivity, affordability, and trust. This SLR therefore establishes the baseline problem that the BLMS is expected to address before the blockchain solution domain is considered (Zein & Twinomurinzi, 2023; Kshetri, 2022).")
    text(doc, "Core contributions of this SLR are:", bold_lead="Core contributions of this SLR are:")
    for item in [
        "A structured specification of Nigerian land-administration failures relevant to secure registration and transfer.",
        "A review of records, identity, document, geospatial, and administrative evidence required for land verification.",
        "A set of evaluation dimensions for measuring any future BLMS improvement honestly.",
        "A socio-technical constraint map that prevents technology from being presented as a substitute for law and institutional governance.",
    ]:
        list_item(doc, item)
    rqs = [
        ["RQ1", "What land-registration and land-administration failures create ownership uncertainty, fraud risk, or delay in Nigeria?", "To identify the core operational problems the BLMS must address."],
        ["RQ2", "What records, data sources, and verification evidence are used to establish or assess land rights?", "To identify the minimum information and evidence model for a trustworthy registry workflow."],
        ["RQ3", "What measures are used to evaluate the effectiveness, efficiency, integrity, and accessibility of land-administration systems?", "To define a credible future evaluation framework for BLMS."],
        ["RQ4", "How do paper-based, centralised digital, and integrated land-registration approaches compare in traceability and service delivery?", "To identify realistic baseline comparisons rather than assume blockchain superiority."],
        ["RQ5", "What legal, institutional, infrastructure, privacy, and inclusion challenges limit land-system improvement in Nigeria?", "To identify implementation constraints and adoption risks."],
    ]
    add_method(doc, rqs, '"land registration", "land administration", "land title", "cadastre", "Nigeria", "property rights", "land records", "e-government"', '(("land registration" OR "land administration" OR "land registry" OR cadastre OR "land title") AND (Nigeria OR "Global South" OR developing) AND (fraud OR transparency OR verification OR digitization OR governance OR "service delivery"))', "January 2018 to July 2026")
    keyword = ASSETS / "problem_keywords.png"
    flow = ASSETS / "problem_flow.png"
    evidence = ASSETS / "problem_evidence.png"
    make_keyword_map(keyword, "Land records and\nverification", "Nigerian land\nadministration", "Fraud, delay and\ntransparency")
    make_workflow(flow)
    make_evidence_map(evidence, [
        ("Record integrity", "Fragmented or weak records make verification costly and uncertain.", "Enforce unique parcel/title identifiers and preserve an audit history."),
        ("Document evidence", "Ownership claims require traceable supporting documents.", "Use content hashing to make document changes detectable."),
        ("Institutional approval", "Legal authority remains with responsible land institutions.", "Keep government approval, correction, and dispute processes explicit."),
        ("Inclusion and trust", "Cost, skills, connectivity, and trust shape adoption.", "Design simple flows and evaluate accessibility before scale-up."),
    ])
    figure(doc, keyword, "Fig. 1. Conceptual search-term map for the BLMS problem-domain review.")
    add_selection_and_quality(doc, flow, evidence,
        ["Published English-language sources from 2018 to July 2026, plus foundational context where directly relevant.", "Sources addressing land registration, land administration, property-right verification, land-record digitisation, or Nigeria/comparable contexts.", "Traceable peer-reviewed articles, conference papers, institutional publications, or publisher records with identifiable authorship."],
        ["Cryptocurrency-only studies with no land, property-right, or public-administration relevance.", "Unverifiable sources, duplicates, promotional claims, and sources that do not state a method, context, or clear contribution.", "Studies that claim legal or performance outcomes without stating the applicable institutional conditions."],
    )
    heading(doc, "3. REVIEW OF PAST LITERATURE ON LAND ADMINISTRATION IN NIGERIA", 2)
    heading(doc, "Traditional Land Registration and Service Delivery", 3)
    text(doc, "Nwuba and Nuhu (2018) examined the challenges of land registration in Kaduna State and highlighted institutional, procedural, and user-facing barriers to effective service delivery. Their study is important because it grounds the BLMS problem in the practical realities of registration rather than treating land administration as a purely technical database problem. Ibrahim et al. (2021) similarly reviewed land-administration weaknesses in Nigeria and linked poor records, limited transparency, and weak administrative processes to the need for improved information management.")
    heading(doc, "Digitisation, Transparency, and Nigerian Land Titles", 3)
    text(doc, "Obamehinti and Eguavoen (2022) reviewed land-title management in Nigeria and argued that transparent, tamper-evident record handling could address weaknesses in conventional title processes. Okoli et al. (2024) further identified inefficiency, corruption risks, unclear procedures, technical constraints, and regulatory challenges as obstacles to modernisation. These sources consistently show that technology adoption must be coupled with staff capacity, legal alignment, and integration with existing land records.")
    heading(doc, "Evidence from Comparable Land-Registration Contexts", 3)
    text(doc, "Zein and Twinomurinzi (2023) synthesised blockchain-related land-registration research in low-income contexts and reported that practical deployment is constrained by institutional resistance, skills, regulation, and infrastructure. Kshetri (2022) likewise cautioned that property-right protection in the Global South requires attention to the political and institutional setting in which digital systems operate. The past literature therefore establishes a clear gap: Nigerian land workflows need better traceability and verification, but a proposed solution must be evaluated as a socio-technical system rather than presented as an automatic legal or anti-fraud remedy.")
    heading(doc, "Land Administration Problems and Traditional Registration Methods", 3)
    text(doc, "The reviewed literature presents land administration as a chain of evidence and decisions rather than a single database operation. A registry must support parcel identification, rights information, supporting documentation, official review, change recording, and public or authorised verification. Failure at any stage can undermine confidence in the result.")
    heading(doc, "RQ1: What land-registration and land-administration failures create ownership uncertainty, fraud risk, or delay in Nigeria?", 3)
    text(doc, "Nigeria-focused studies describe administrative delay, inadequate record management, cost, unclear procedures, limited staff capacity, and low transparency as recurring concerns (Nwuba & Nuhu, 2018; Ibrahim et al., 2021; Okoli et al., 2024). These conditions can make it difficult for a purchaser, surveyor, lender, or land authority to establish a current and consistent view of a parcel and its claimed owner. Multiple sale or allocation risks arise when the same land can be represented through inconsistent records or when verification is slow and discretionary. The problem is therefore a combination of process design, source-data integrity, and institutional accountability rather than simply the absence of a blockchain.")
    heading(doc, "RQ2: What records, data sources, and verification evidence are used to establish or assess land rights?", 3)
    text(doc, "Land verification typically combines descriptive parcel identifiers, survey or geospatial references, title or allocation information, conveyance documents, identity evidence, payment/levy records, and the relevant authority's decisions. The literature also distinguishes between the data held in a registry and the documents that support a transaction (Obamehinti & Eguavoen, 2022; Zein & Twinomurinzi, 2023). A system should not expose all personal or evidential data indiscriminately; it needs appropriate access control, retention, and correction procedures. For BLMS, this implies a controlled data model with parcel and title uniqueness, role-separated actions, and document-content verification rather than a claim that a file upload proves legal title.")
    heading(doc, "RQ3: What measures are used to evaluate land-administration systems?", 3)
    text(doc, "The literature suggests that system evaluation must combine service and governance measures. Transaction time, cost, error or duplication rate, search/verification time, record completeness, auditability, user trust, accessibility, privacy protection, and dispute/correction handling are all more meaningful than a technology label (Zein & Twinomurinzi, 2023; Okoli et al., 2024). The following table converts those measures into a future BLMS evaluation frame.")
    text(doc, "Table 2. Evaluation dimensions for land-administration systems.", bold_lead="Table 2.")
    table(doc, ["Dimension", "Example measure", "Interpretation for BLMS"], [
        ["Efficiency", "Time from submission to authorised registration; search time; user cost.", "Measure workflow improvement against a declared centralised or paper-process baseline."],
        ["Integrity", "Duplicate parcel/title attempts rejected; document MATCH/MISMATCH outcomes; audit completeness.", "Shows whether the system enforces and records stated business rules."],
        ["Transparency", "Availability of authorised history and decision trail; ability to explain status.", "Shows traceability, not public exposure of sensitive personal data."],
        ["Accessibility", "Task completion, error rate, usability feedback, support needs, digital-access constraints.", "Prevents a technically secure platform from excluding intended users."],
        ["Governance", "Authorisation correctness, correction path, dispute handling, privacy compliance.", "Confirms that technology remains accountable to lawful institutions."],
    ], [1760, 3600, 4000])
    heading(doc, "RQ4: How do paper-based, centralised digital, and integrated approaches compare?", 3)
    text(doc, "Paper-heavy processes can make records difficult to search, compare, and preserve consistently, while centralised digitisation may improve retrieval but still concentrate control and leave unclear audit trails if governance is weak. An integrated approach can improve shared visibility and structured workflows, but its benefits depend on trustworthy input data and accountable institutions (Kshetri, 2022; Obamehinti & Eguavoen, 2022). The relevant comparison is therefore not an absolute claim that one architecture is always superior; it is whether the selected process demonstrably reduces verification effort, duplicate risk, and unexplained changes under stated operating conditions.")
    heading(doc, "RQ5: What legal, institutional, infrastructure, privacy, and inclusion challenges limit land-system improvement in Nigeria?", 3)
    text(doc, "The evidence identifies legal alignment, existing records, institutional coordination, staff capacity, infrastructure, funding, privacy, and user trust as material constraints (Ibrahim et al., 2021; Okoli et al., 2024; Zein & Twinomurinzi, 2023). Government institutions remain responsible for legal recognition, adjudication, and correction. Digital systems must also accommodate unequal connectivity and digital literacy, protect personal data, and avoid presenting a prototype credential or synthetic identity as a legal identity proof. These constraints are design requirements for the BLMS, not matters to postpone until after deployment.")
    heading(doc, "CONCLUSION", 3)
    text(doc, "This problem-domain SLR shows that the BLMS is justified by persistent needs for better traceability, faster and clearer verification, controlled document integrity, role accountability, and more transparent land-service workflows. It also shows that the project cannot promise legal title, fraud elimination, or nationwide performance solely through software. The system must be evaluated against service, integrity, governance, inclusion, and legal criteria. These findings define the baseline against which the separate solution-domain SLR assesses blockchain, Hyperledger Fabric, smart contracts, and content-addressed documents.")
    add_references(doc, [
        "Ibrahim, I., Daud, D., Azmi, F. A. M., Noor, N. A. M., & Yusoff, N. S. M. (2021). Improvement of land administration system in Nigeria: A blockchain technology review. International Journal of Scientific & Technology Research, 10(8), 33-39.",
        "Nwuba, C. C., & Nuhu, S. R. (2018). Challenges to land registration in Kaduna State, Nigeria. Journal of African Real Estate Research, 3(1), 141-172. https://doi.org/10.15641/jarer.v3i1.546",
        "Obamehinti, A. S., & Eguavoen, V. (2022). A literature review of land title with the aim of maximising the benefits of blockchain technology in the management of land title in Nigeria. Studia Universitatis Babes-Bolyai Engineering, 67(1), 124-135. https://doi.org/10.24193/subbeng.2022.1.12",
        "Okoli, F. U., Oludiji, S. M., Ofoegbunam, E. I., Oyesiji, O. M., & Akindiya, O. M. (2024). Blockchain technology for land registration in Nigeria: A review of opportunities and challenges. FUDMA Journal of Sciences, 8(6). https://doi.org/10.33003/fjs-2024-0806-2919",
        "Zein, R. M., & Twinomurinzi, H. (2023). Blockchain technology in lands registration: A systematic literature review. Journal of eDemocracy and Open Government, 15(2), 1-36. https://doi.org/10.29379/jedem.v15i2.748",
        "Kshetri, N. (2022). Blockchain as a tool to facilitate property rights protection in the Global South: Lessons from India. Journal of Information Technology, 37(4), 391-409.",
        "Page, M. J., McKenzie, J. E., Bossuyt, P. M., et al. (2021). The PRISMA 2020 statement: An updated guideline for reporting systematic reviews. BMJ, 372, n71. https://doi.org/10.1136/bmj.n71",
        "Kitchenham, B., & Charters, S. (2007). Guidelines for performing systematic literature reviews in software engineering. EBSE Technical Report, EBSE-2007-01.",
    ])
    path = OUT / "BLMS_SLR_Problem_Domain.docx"
    doc.save(path)
    return path


def build_solution():
    doc = doc_base(1.25)
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("A Systematic Literature Review on the Solution Domain of a Blockchain-Based Land Management System")
    set_run(r, 18, bold=True)
    text(doc, "ABSTRACT", bold_lead="ABSTRACT")
    text(doc, "This systematic literature review examines the solution domain for the Blockchain-Based Land Management System (BLMS): permissioned blockchain architecture, Hyperledger Fabric, Go chaincode, smart contracts, content-addressed document verification, identity, privacy, and evaluation. The review follows a structured protocol and addresses five questions about blockchain architectures for land registration, relevant data and evidence models, evaluation metrics, comparisons with centralised approaches, and adoption constraints. The evidence supports blockchain most strongly as a controlled shared audit and coordination layer where several authorised parties need a consistent transaction history. Hyperledger Fabric is especially relevant to a citizen-government model because it provides permissioned membership, organisational identities, endorsement policies, and programmable contract rules. Content-addressed document checks can make file changes detectable without storing large documents on-chain. The review also finds that blockchain does not validate false source data, remove legal authority, guarantee performance, or eliminate governance and privacy obligations. The resulting implications support a phased BLMS prototype and a separate, measurable Fabric-integration phase.")
    text(doc, "Keywords: blockchain; Hyperledger Fabric; smart contracts; Go chaincode; IPFS; content addressing; land registration; digital identity.", bold_lead="Keywords:")
    heading(doc, "1. INTRODUCTION", 1)
    text(doc, "The solution domain addresses the technology and governance mechanisms that may improve the land-administration problems identified in the companion review. A land registry is not improved merely by copying records to a ledger. It requires controlled identities, defined roles, transaction rules, privacy safeguards, documentary evidence, institutional approval, and a way to investigate or correct disputed entries. Blockchain must therefore be assessed as one architectural component in a wider socio-technical system (Batubara et al., 2018; Zein & Twinomurinzi, 2023).")
    text(doc, "BLMS proposes a permissioned design in which citizen and government actors have distinct roles. The local prototype demonstrates the workflow without requiring Docker or an external service, while the repository contains Go chaincode as a future Fabric research artifact. This SLR explains what those technologies can contribute, what they cannot prove, and how future evaluation should be designed (Androulaki et al., 2018; Zein & Twinomurinzi, 2024).")
    text(doc, "Core contributions of this SLR are:", bold_lead="Core contributions of this SLR are:")
    for item in [
        "A comparison of blockchain, permissioned-ledger, smart-contract, identity, and document-integrity roles in land registration.",
        "A BLMS-specific interpretation of why Hyperledger Fabric and Go chaincode are appropriate for a two-organisation research architecture.",
        "A clear evaluation frame for later Fabric tests without claiming benchmarks that have not been run.",
        "An adoption and governance map covering legal authority, privacy, interoperability, skills, infrastructure, and data quality.",
    ]:
        list_item(doc, item)
    rqs = [
        ["RQ1", "How are blockchain and permissioned-ledger architectures applied to land registration and property-right workflows?", "To identify the architectural patterns relevant to BLMS."],
        ["RQ2", "What data, identity, document, and geospatial evidence models are used in blockchain-enabled land systems?", "To define which information should be on-chain, off-chain, or content-addressed."],
        ["RQ3", "What metrics are used to evaluate blockchain land-registration and Hyperledger Fabric systems?", "To define honest future performance and correctness evaluation."],
        ["RQ4", "How do blockchain-enabled land systems compare with centralised or conventional approaches?", "To assess reported benefits and trade-offs without assuming superiority."],
        ["RQ5", "What infrastructure, governance, legal, privacy, and adoption challenges constrain deployment?", "To identify the conditions required before a Fabric-connected BLMS can be claimed as operational."],
    ]
    add_method(doc, rqs, '"blockchain", "Hyperledger Fabric", "land registration", "land registry", "smart contract", "IPFS", "content addressing", "digital identity"', '(("land registration" OR "land registry" OR "land title" OR "property registration") AND (blockchain OR "distributed ledger" OR "Hyperledger Fabric" OR "smart contract" OR IPFS) AND (identity OR document OR performance OR governance OR privacy OR adoption))', "January 2018 to July 2026")
    keyword = ASSETS / "solution_keywords.png"
    flow = ASSETS / "solution_flow.png"
    evidence = ASSETS / "solution_evidence.png"
    make_keyword_map(keyword, "Permissioned ledger\nand chaincode", "Blockchain land\nmanagement", "Identity, documents\nand privacy")
    make_workflow(flow)
    make_evidence_map(evidence, [
        ("Permissioned governance", "Fabric supports known organisations, roles, and endorsement policy.", "Model distinct citizen and government permissions."),
        ("Transaction rules", "Smart contracts enforce repeatable state-transition checks.", "Use Go chaincode for duplicate, role, listing, levy, and transfer rules."),
        ("Off-chain documents", "Content identifiers can reveal whether the file has changed.", "Store/compare document hashes; do not claim live IPFS unless deployed."),
        ("Evaluation and risk", "Performance and trust depend on network design, data, law, and operating conditions.", "Run stated Fabric and usability tests before making operational claims."),
    ])
    figure(doc, keyword, "Fig. 1. Conceptual search-term map for the BLMS solution-domain review.")
    add_selection_and_quality(doc, flow, evidence,
        ["Published English-language work from 2018 to July 2026, plus foundational Fabric or review-method work directly needed for the RQs.", "Sources that address land/property registration, permissioned or public blockchain, smart contracts, document integrity, identity, or blockchain evaluation.", "Traceable journal articles, conference papers, recognised technical publications, and publisher or institutional records."],
        ["Blockchain studies unrelated to land/property workflows or public administration.", "Unverifiable repositories, duplicated records, marketing material, and sources without an identifiable method, architecture, or limitation.", "Claims of legal transfer, real identity verification, or benchmark performance without stated conditions and evidence."],
    )
    heading(doc, "3. REVIEW OF PAST LITERATURE ON BLOCKCHAIN-BASED LAND MANAGEMENT", 2)
    heading(doc, "Blockchain Land-Registry Proposals", 3)
    text(doc, "Early land-registry proposals established the main technical idea: represent a parcel or title as a digital asset, record changes as transactions, and use cryptographic links to produce an auditable history. Gupta et al. (2019) proposed LandLedger for land-property administration, while Krishnapriya and Sarath (2020) and Biswas et al. (2021) presented blockchain-based approaches to secure land registration. Alam et al. (2022) developed a land-title management design for Bangladesh. Together, these studies show that blockchain can support traceability and rule-based workflow, but most are architecture proposals or prototypes rather than evidence of nationwide legal deployment.")
    heading(doc, "Permissioned Blockchain and Hyperledger Fabric", 3)
    text(doc, "Androulaki et al. (2018) described Hyperledger Fabric as a permissioned blockchain platform with organisational identities, endorsement policies, and programmable chaincode. These features make Fabric relevant to public-service workflows in which participants are known and responsibilities must be separated. Zein and Twinomurinzi (2024) applied Fabric to information sharing in land registration, demonstrating why a permissioned design is more appropriate than an unrestricted public network when land agencies must control membership, data access, and approval rules.")
    heading(doc, "Document Integrity, Identity, and Adoption Literature", 3)
    text(doc, "Banerjee et al. (2022) combined blockchain and IPFS for a reliable land-registry design, illustrating the use of content identifiers to make document changes detectable. Liu et al. (2020) reviewed blockchain-based identity management and highlighted the importance of identity, privacy, and credential governance. Batubara et al. (2018) reviewed e-government adoption challenges and found that governance, legal frameworks, skills, and infrastructure can be as significant as the technical platform. This literature gives BLMS a concrete solution direction but also identifies a gap: a locally demonstrable land workflow must be followed by a separately evaluated permissioned-network deployment rather than assuming that source code alone proves production suitability.")
    heading(doc, "Blockchain-Enabled Land Registration and Solution Architectures", 3)
    text(doc, "The solution literature presents several recurring patterns: a ledger records land-state changes and transactions; contracts check business rules; documents and geospatial data remain outside the ledger or are represented by hashes; identities and roles control who can propose or approve changes; and an authority or consortium governs the network. These patterns differ in whether they use public chains, private chains, or permissioned consortium platforms (Alam et al., 2022; Gupta et al., 2019; Krishnapriya & Sarath, 2020). For a land-administration context, the main design question is how to balance integrity, privacy, legal authority, cost, and operational feasibility.")
    heading(doc, "RQ1: How are blockchain and permissioned-ledger architectures applied to land registration and property-right workflows?", 3)
    text(doc, "Land-registry proposals commonly model a parcel or title as an asset and represent registration, transfer, verification, and update as transactions. Permissioned systems are attractive where the participants are known public or institutional actors because membership, roles, and transaction approval can be governed. Hyperledger Fabric provides a modular permissioned architecture with identities, channels or private data options, endorsement policies, and chaincode (Androulaki et al., 2018; Zein & Twinomurinzi, 2024). These features align with a two-organisation BLMS research model in which citizen and government actors have different permissions. They do not by themselves establish legal title or make a bad source record correct.")
    heading(doc, "RQ2: What data, identity, document, and geospatial evidence models are used?", 3)
    text(doc, "The literature generally advises against placing every document and personal data item directly on a blockchain. A practical model separates land-state metadata from supporting files and sensitive identity information. The ledger can hold a parcel identifier, title reference, owner reference or pseudonymous key, status, timestamps, transaction history, and a content identifier for an associated document (Banerjee et al., 2022; Liu et al., 2020). IPFS is one option for content-addressed storage; the core integrity mechanism is that the identifier changes when the file content changes. Geospatial coordinates and survey data may be recorded as structured references subject to validation. Privacy and access rules remain essential, especially where a land record links identity, location, and economic information.")
    heading(doc, "RQ3: What metrics are used to evaluate blockchain land-registration and Hyperledger Fabric systems?", 3)
    text(doc, "Blockchain evaluation must include both software performance and business correctness. Throughput, latency, success rate, query response time, resource use, scalability, and fault behaviour are common systems metrics (Thakkar et al., 2018; Wang & Chu, 2020). For a land workflow, duplicate prevention, authorisation correctness, history availability, document-verification correctness, and error handling are equally important. Comparisons require a declared workload, network topology, hardware, payload size, endorsement policy, data volume, and baseline. Without those conditions, a transaction-per-second statement is not transferable to BLMS.")
    text(doc, "Table 2. Evaluation metrics for a future Fabric-connected BLMS.", bold_lead="Table 2.")
    table(doc, ["Metric family", "Illustrative measure", "Interpretation"], [
        ["Ledger performance", "Submit/commit latency, throughput, failed transaction rate, query response time.", "Report with workload, organisation/peer topology, hardware, payload, and policy."],
        ["Functional correctness", "Duplicate attempts rejected; role violations rejected; valid workflow completion.", "Confirms chaincode and application rules behave as specified."],
        ["Integrity and audit", "History retrieval, event delivery, document MATCH/MISMATCH accuracy.", "Shows traceability and evidence consistency, not legal finality."],
        ["Security and privacy", "Authorisation test results, sensitive-data exposure, key/identity lifecycle coverage.", "Tests controls appropriate to a permissioned public-service system."],
        ["Adoption and service", "Task success, user error, training/support needs, cost, accessibility feedback.", "Captures whether technical design is usable in its operational context."],
    ], [1760, 3600, 4000])
    heading(doc, "RQ4: How do blockchain-enabled land systems compare with centralised or conventional approaches?", 3)
    text(doc, "Studies often report improved auditability, traceability, and resistance to unauthorised alteration when a distributed or permissioned ledger is used (Zein & Twinomurinzi, 2023; Khalid et al., 2022). These are plausible design advantages where more than one authorised organisation must reconcile changes. At the same time, a centralised registry can be faster, simpler, and cheaper to operate when it already has effective governance and high-quality data. The strongest defensible comparison is therefore conditional: blockchain may add value where shared control, independent verification, and a consistent transaction history are genuine requirements. It imposes additional complexity in governance, certificates, policy configuration, monitoring, and integration.")
    heading(doc, "RQ5: What infrastructure, governance, legal, privacy, and adoption challenges constrain deployment?", 3)
    text(doc, "The most important deployment barriers are rarely solved by chaincode alone. They include legal recognition of electronic processes, the handling of legacy records, institutional ownership of nodes and keys, identity assurance, data-protection obligations, network availability, operating skills, procurement cost, interoperability, and clear correction/dispute procedures (Batubara et al., 2018; Liu et al., 2020; Zein & Twinomurinzi, 2024). The local BLMS prototype intentionally avoids these dependencies so it can be tested without Docker or external services. A later Fabric phase must explicitly introduce and evaluate these dependencies rather than treating them as invisible implementation details.")
    heading(doc, "CONCLUSION", 3)
    text(doc, "This solution-domain SLR supports a permissioned, phased BLMS architecture. Hyperledger Fabric is appropriate for the research direction because it can encode organisational membership, role-aware policy, and deterministic land-workflow rules in Go chaincode. Content-addressed documents can support integrity checks without putting full files on the ledger. The evidence also establishes clear limits: a blockchain cannot validate false inputs, replace statutory decision-making, establish real identity without an authorised integration, or prove performance without a measured workload. The current local prototype is therefore a valid workflow demonstration; the next research result must be a transparent Fabric deployment and evaluation under stated conditions.")
    add_references(doc, [
        "Alam, K. M., Rahman, J. A., Tasnim, A., & Akther, A. (2022). A blockchain-based land title management system for Bangladesh. Journal of King Saud University - Computer and Information Sciences, 34(6), 3096-3110. https://doi.org/10.1016/j.jksuci.2020.10.011",
        "Androulaki, E., Barger, A., Bortnikov, V., et al. (2018). Hyperledger Fabric: A distributed operating system for permissioned blockchains. Proceedings of the Thirteenth EuroSys Conference, 1-15. https://doi.org/10.1145/3190508.3190538",
        "Banerjee, S., Kumar, K., Masulkar, P., Amin, R., & Dwivedi, S. K. (2022). Blockchain and IPFS-based reliable land registry system. Security and Privacy, 5(5), e236. https://doi.org/10.1002/spy2.236",
        "Batubara, F. R., Ubacht, J., & Janssen, M. (2018). Challenges of blockchain technology adoption for e-government: A systematic literature review. Proceedings of the 19th Annual International Conference on Digital Government Research, 1-9. https://doi.org/10.1145/3209281.3209317",
        "Biswas, M., Al Faysal, J., & Ahmed, K. A. (2021). LandChain: A blockchain-based secured land registration system. 2021 International Conference on Science & Contemporary Technologies, 1-6. https://doi.org/10.1109/ICSCT53883.2021.9642669",
        "Gupta, N., Das, M. L., & Nandi, S. (2019). LandLedger: Blockchain-powered land property administration system. 2019 IEEE International Conference on Advanced Networks and Telecommunications Systems, 1-6. https://doi.org/10.1109/ANTS47819.2019.9118071",
        "Khalid, M. I., Iqbal, J., Alturki, A., Hussain, S., Alabrah, A., & Ullah, S. S. (2022). Blockchain-based land registration system: A conceptual framework. Applied Bionics and Biomechanics, 2022, 3859629. https://doi.org/10.1155/2022/3859629",
        "Krishnapriya, S., & Sarath, G. (2020). Securing land registration using blockchain. Procedia Computer Science, 171, 1708-1715. https://doi.org/10.1016/j.procs.2020.04.183",
        "Liu, Y., He, D., Obaidat, M. S., Kumar, N., Khan, M. K., & Choo, K.-K. R. (2020). Blockchain-based identity management systems: A review. Journal of Network and Computer Applications, 166, 102731. https://doi.org/10.1016/j.jnca.2020.102731",
        "Thakkar, P., Nathan, S., & Viswanathan, B. (2018). Performance benchmarking and optimizing Hyperledger Fabric blockchain platform. 2018 IEEE 26th International Symposium on Modeling, Analysis, and Simulation of Computer and Telecommunication Systems, 264-276. https://doi.org/10.1109/MASCOTS.2018.00034",
        "Wang, C., & Chu, X. (2020). Performance characterization and bottleneck analysis of Hyperledger Fabric. 2020 IEEE 40th International Conference on Distributed Computing Systems, 1281-1286. https://doi.org/10.1109/ICDCS47774.2020.00165",
        "Zein, R. M., & Twinomurinzi, H. (2024). Information sharing in land registration using Hyperledger Fabric blockchain. Digital, 2(2), 6. https://doi.org/10.3390/digital2020006",
    ])
    path = OUT / "BLMS_SLR_Solution_Domain.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_problem())
    print(build_solution())
