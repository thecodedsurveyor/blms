from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
OUT.mkdir(exist_ok=True)

BLUE = "1F4D78"
MID_BLUE = "2E74B5"
INK = "1F2937"
MUTED = "5B6573"
PALE = "E8EEF5"
PALE_2 = "F4F6F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = 120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = tc_pr.first_child_found_in("w:tcMar")
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for name, value in (("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")):
                el = mar.find(qn(f"w:{name}"))
                if el is None:
                    el = OxmlElement(f"w:{name}")
                    mar.append(el)
                el.set(qn("w:w"), value)
                el.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, *, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, field_code: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def base_document(running_label: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, MID_BLUE, 18, 10),
        ("Heading 2", 13, MID_BLUE, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Table Text" not in [s.name for s in doc.styles]:
        table_style = doc.styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
        table_style.font.name = "Calibri"
        table_style.font.size = Pt(9.5)
        table_style.paragraph_format.space_after = Pt(0)
        table_style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run(running_label.upper())
    set_run_font(r, size=9, color=MUTED, bold=True)
    header.paragraph_format.space_after = Pt(0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("BLMS | Page ")
    set_run_font(r, size=9, color=MUTED)
    add_field(footer, "PAGE")
    return doc


def paragraph(doc: Document, text: str = "", *, bold_lead: str | None = None, italic: bool = False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r)
    return p


def number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_run_font(r)
    return p


def title_page(doc: Document, title: str, subtitle: str, doc_type: str):
    for _ in range(7):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("BLOCKCHAIN-BASED LAND MANAGEMENT SYSTEM (BLMS)")
    set_run_font(r, size=12, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(title)
    set_run_font(r, size=23, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run(subtitle)
    set_run_font(r, size=13, color=MUTED, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(doc_type)
    set_run_font(r, size=12, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(86)
    r = p.add_run("Prepared by\nOMO LOUIS\nDepartment of Computer Science\nFaculty of Science")
    set_run_font(r, size=12, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(52)
    r = p.add_run("July 2026")
    set_run_font(r, size=11, color=MUTED)
    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, PALE)
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Text"]
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Text"]
            r = p.add_run(text)
            set_run_font(r, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc: Document, label: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_2)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " ")
    set_run_font(r, size=10.5, color=BLUE, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_references(doc: Document, refs: list[str]):
    doc.add_heading("References", level=1)
    for item in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        set_run_font(r, size=10)


def build_slr():
    doc = base_document("Systematic Literature Review")
    title_page(
        doc,
        "Systematic Literature Review",
        "Problem and solution domains for a blockchain-based land management system in Nigeria",
        "Lecturer Submission",
    )
    doc.add_heading("Abstract", level=1)
    paragraph(doc, "This review examines the problem and solution domains relevant to the proposed Blockchain-Based Land Management System (BLMS) for Nigeria. It synthesises peer-reviewed and scholarly work on land-administration failures, blockchain-enabled land registration, permissioned ledgers, content-addressed document verification, smart contracts, and adoption barriers. A protocol-led search was undertaken on 18 July 2026 using Google Scholar-compatible web discovery, publisher pages, institutional repositories, and DOI records. The eligible evidence consistently identifies weak record integrity, opaque verification, fragmented institutions, delays, and double allocation risks as central problem-domain concerns. The solution literature supports blockchain principally as a shared audit and coordination layer, not as a substitute for legal adjudication or reliable source data. Hyperledger Fabric is relevant because permissioned membership, organisational separation, and endorsement can reflect a citizen-government setting; content addressing can make document changes detectable without placing documents directly on-chain. The review therefore supports a phased BLMS design, beginning with demonstrable workflow integrity and explicitly reserving live identity, IPFS, Fabric-network deployment, performance benchmarking, and legal integration for separately evaluated phases.")
    paragraph(doc, "Keywords: land administration; land registration; Nigeria; blockchain; Hyperledger Fabric; IPFS; smart contracts; systematic literature review.", bold_lead="Keywords:")

    doc.add_heading("1. Introduction", level=1)
    paragraph(doc, "Land administration depends on trustworthy evidence of rights, parcels, transfers, and official decisions. In Nigeria, the literature repeatedly associates current processes with slow registration, weak transparency, fragmented records, multiple sales or allocations, cost, and institutional bottlenecks. The project question is therefore not merely whether a blockchain can store a title. It is how a secure digital workflow can improve traceability and accountability while respecting the continuing legal and institutional role of land authorities.")
    paragraph(doc, "This document is a structured systematic literature review and a reproducible protocol, prepared for the BLMS project. It follows the reporting intent of PRISMA 2020 and the protocol discipline recommended for software-engineering SLRs by Kitchenham and Charters. It does not invent database counts or claim exhaustive access to subscription databases. The included evidence list and screening logic are supplied so a supervisor can reproduce or extend the review through Scopus, Web of Science, IEEE Xplore, or institutional-library access.")

    doc.add_heading("2. Review protocol", level=1)
    doc.add_heading("2.1 Review objective and research questions", level=2)
    paragraph(doc, "The objective is to synthesise evidence on the land-administration problem domain and the technical, legal, and organisational solution domain of blockchain-enabled land registration, with a Nigeria-focused interpretation for the BLMS.")
    add_table(doc, ["ID", "Research question", "Use in BLMS"], [
        ["RQ1", "What land-administration failures are reported in Nigeria and comparable low-income contexts?", "Defines the problem and user requirements."],
        ["RQ2", "What benefits and limitations are reported for blockchain-enabled land registration?", "Tests whether blockchain is justified rather than assumed."],
        ["RQ3", "Which architectural roles are reported for permissioned blockchains, smart contracts, and content-addressed documents?", "Guides the Fabric, Go-chaincode, and document-verification design."],
        ["RQ4", "What adoption, legal, identity, governance, and infrastructure barriers remain?", "Defines non-technical requirements and risk controls."],
        ["RQ5", "What evidence gaps must be addressed by the BLMS implementation and its evaluation?", "Separates implemented evidence from future research claims."],
    ], [720, 3960, 4680])

    doc.add_heading("2.2 Search strategy", level=2)
    paragraph(doc, "Searches were conducted on 18 July 2026. Discovery routes were publisher pages, DOI records, institutional repositories, and scholarly web indexes. The date, search expressions, inclusion rules, and source list are retained below to make the process auditable. A future dissertation version should repeat these strings in the university's subscribed databases and record the returned counts before claiming an exhaustive PRISMA flow.")
    add_table(doc, ["Concept", "Representative search expression"], [
        ["Land and blockchain", '("land registration" OR "land registry" OR cadastre OR "land administration") AND (blockchain OR "distributed ledger")'],
        ["Architecture", '("land registration" OR "property registration") AND ("Hyperledger Fabric" OR permissioned OR "smart contract")'],
        ["Documents", '("land registry" OR "property title") AND (IPFS OR "content addressed" OR "document integrity")'],
        ["Nigeria focus", '(Nigeria OR "Global South" OR "low-income countries") AND ("land administration" OR "land title") AND blockchain'],
    ], [2520, 6840])
    doc.add_heading("2.3 Eligibility and quality appraisal", level=2)
    add_table(doc, ["Criterion", "Decision rule"], [
        ["Time and language", "English-language scholarly sources published from 2018 to 18 July 2026; foundational Fabric and review-method sources retained where directly relevant."],
        ["Topical relevance", "Must address land registration/administration, blockchain or a directly relevant architecture, or Nigeria/Global-South land-administration constraints."],
        ["Evidence quality", "Peer-reviewed articles, recognised conference papers, authoritative method guidance, and institutional repositories. Blog posts and unverifiable claims were excluded."],
        ["Exclusions", "Cryptocurrency-only studies, studies without a land-registration connection, duplicate records, and claims without identifiable authorship or publication details."],
        ["Appraisal questions", "Is the context clear? Is the method or architecture described? Are limitations stated? Is the source traceable through a DOI, publisher, or institutional record?"],
    ], [2160, 7200])
    add_callout(doc, "Transparency note.", "The review is systematic in protocol and traceability, but it is not represented as a database-exhaustive meta-analysis. No unverified record counts, effect sizes, or PRISMA totals are reported.")

    doc.add_heading("3. Results: problem domain", level=1)
    doc.add_heading("3.1 Record integrity, verification, and transaction risk", level=2)
    paragraph(doc, "Nigeria-focused reviews identify delay, opacity, inadequate technical infrastructure, unclear procedures, high transaction costs, limited skills, weak record management, and opportunities for multiple sales or unauthorised alteration as recurring concerns. The practical consequence is that a purchaser or agency must repeatedly verify whether a parcel, title, seller, supporting document, and payment instruction can be trusted. These are workflow and governance problems as much as database problems.")
    doc.add_heading("3.2 Fragmented authority and institutional coordination", level=2)
    paragraph(doc, "Land rights derive from law and administrative decisions. A digital ledger cannot itself validate a defective survey, settle a dispute, or replace statutory approval. The literature therefore cautions against treating decentralisation as a cure for poor source data. Its more defensible role is to give authorised participants a shared, tamper-evident transaction history while preserving the authority of land agencies and adjudication processes.")
    doc.add_heading("3.3 Trust and inclusion", level=2)
    paragraph(doc, "Trustworthy registration also depends on accessible procedures, identity assurance, privacy, support for citizens who have limited digital access, and institutional legitimacy. The studies on low-income countries report that government resistance, skills, regulation, and infrastructure can constrain real deployment even when the technical concept is appealing. This shifts the BLMS emphasis from a generic 'blockchain solution' to a staged socio-technical system.")

    doc.add_heading("4. Results: solution domain", level=1)
    doc.add_heading("4.1 Permissioned blockchain and Hyperledger Fabric", level=2)
    paragraph(doc, "Hyperledger Fabric is a permissioned distributed-ledger architecture designed around identifiable organisations, policies, and modular components. For a land-management setting, these characteristics can support a two-organisation model in which citizen-facing and government roles are represented separately, while endorsement policies determine whose approval is required for sensitive changes. Fabric is therefore more aligned with administrative accountability than an unrestricted public chain. Nevertheless, Fabric does not remove the need for operating governance, key management, network administration, and legal rules.")
    doc.add_heading("4.2 Smart contracts and Go chaincode", level=2)
    paragraph(doc, "Land workflows can be expressed as controlled state transitions: register a parcel, reject duplicate identifiers, attach or update evidence, list a property, record an approved transfer, update a levy, and preserve history. Chaincode or smart contracts provide a place to apply these preconditions consistently. Their value is deterministic enforcement and auditable history; they must not be described as automatic legal conveyance unless the relevant law and agency processes explicitly make them so.")
    doc.add_heading("4.3 Content-addressed document verification", level=2)
    paragraph(doc, "A document-integrity design computes a cryptographic digest from an uploaded title or supporting document and records the resulting content identifier. Recomputing the identifier later allows a MATCH or MISMATCH result, making changes to the file detectable. This can improve evidential traceability while avoiding storage of large documents directly on the ledger. IPFS is one possible content-addressed storage network; in the current BLMS local prototype the same integrity idea is demonstrated entirely in the browser, without claiming a live IPFS service.")
    doc.add_heading("4.4 Identity, privacy, and governance", level=2)
    paragraph(doc, "Identity links real-world people to digital actions, but it also creates privacy and security obligations. The literature favours careful identity integration, role separation, data minimisation, consent, and auditability. In the BLMS project, NIN behaviour is synthetic only. Live National Identity Management Commission integration is not implemented and must not be inferred from a prototype login or test record.")

    doc.add_heading("5. Synthesis and implications for BLMS", level=1)
    add_table(doc, ["Evidence theme", "Design implication", "Current BLMS position"], [
        ["Duplicate and opaque records", "Use unique parcel/title rules, role checks, auditable events, and clear history.", "Implemented locally in the browser adapter and optional API mock ledger."],
        ["Document tampering concern", "Use content hashing and compare candidate files against the recorded identifier.", "Implemented locally; no live IPFS node is claimed."],
        ["Institutional accountability", "Use roles, approval rules, and a permissioned-network design for the Fabric research track.", "Go chaincode source exists; a Fabric network and Gateway integration remain future work."],
        ["Adoption and legal barriers", "Treat law, agency workflow, identity governance, training, and infrastructure as first-class requirements.", "Documented as constraints; not solved by the prototype."],
        ["Need for empirical evidence", "Evaluate usability, correctness, resilience, throughput, cost, and governance with a declared protocol.", "No Caliper benchmark or 15-person expert evaluation has been executed or reported."],
    ], [2340, 3900, 3120])
    paragraph(doc, "The synthesis supports BLMS as a phased research prototype. The minimum viable contribution is a traceable local workflow that demonstrates the problem-solution mapping honestly. The next research contribution is not simply 'turning Fabric on'; it is conducting a controlled Fabric deployment and measuring whether it improves the required properties under stated conditions.")

    doc.add_heading("6. Threats to validity and future review work", level=1)
    bullet(doc, "Coverage risk: this review uses traceable web-accessible scholarly sources, but should be extended through library database searches before final thesis submission.")
    bullet(doc, "Publication bias: proposal papers may state benefits more strongly than field deployments demonstrate; implementation claims should be distinguished from conceptual designs.")
    bullet(doc, "Context transfer: international land-registry examples are informative but do not automatically transfer to Nigeria's legal, institutional, or infrastructure conditions.")
    bullet(doc, "Terminology: 'immutable' describes record-evidence properties, not an unconditional legal outcome. Authorised correction, dispute resolution, and data-protection procedures must be designed explicitly.")
    bullet(doc, "Future update: repeat the search in Scopus, Web of Science, IEEE Xplore, and Google Scholar; export records; deduplicate; screen independently where possible; then publish an exact PRISMA flow diagram.")

    add_references(doc, [
        "Androulaki, E., Barger, A., Bortnikov, V., et al. (2018). Hyperledger Fabric: A distributed operating system for permissioned blockchains. Proceedings of the Thirteenth EuroSys Conference. https://doi.org/10.1145/3190508.3190538",
        "Banerjee, M., Lee, J., & Choo, K.-K. R. (2022). Blockchain and IPFS-based reliable land registry system. Security and Privacy, 5(2), e236. https://doi.org/10.1002/spy2.236",
        "Khalid, M. I., Iqbal, J., Hussain, A. A., & Ullah, S. S. (2022). Blockchain-based land registration system: A conceptual framework. Applied Bionics and Biomechanics, 2022, 3859629. https://doi.org/10.1155/2022/3859629",
        "Kitchenham, B., & Charters, S. (2007). Guidelines for performing systematic literature reviews in software engineering. EBSE Technical Report, EBSE-2007-01.",
        "Krishnapriya, S., & Sarath, G. (2020). Securing land registration using blockchain. Procedia Computer Science, 171, 1708-1715. https://doi.org/10.1016/j.procs.2020.04.183",
        "Obamehinti, A. S., & Eguavoen, V. (2022). A literature review of land title with the aim of maximising the benefits of blockchain technology in the management of land title in Nigeria. Studia Universitatis Babes-Bolyai Engineering, 67(1), 124-135. https://doi.org/10.24193/subbeng.2022.1.12",
        "Okoli, F. U., Oludiji, S. M., Ofoegbunam, E. I., Oyesiji, O. M., & Akindiya, O. M. (2024). Blockchain technology for land registration in Nigeria: A review of opportunities and challenges. FUDMA Journal of Sciences, 8(6). https://doi.org/10.33003/fjs-2024-0806-2919",
        "Page, M. J., McKenzie, J. E., Bossuyt, P. M., et al. (2021). The PRISMA 2020 statement: An updated guideline for reporting systematic reviews. BMJ, 372, n71. https://doi.org/10.1136/bmj.n71",
        "Zein, R. M., & Twinomurinzi, H. (2023). Blockchain technology in lands registration: A systematic literature review. Journal of eDemocracy and Open Government, 15(2), 1-36. https://doi.org/10.29379/jedem.v15i2.748",
        "Zein, R. M., & Twinomurinzi, H. (2024). Information sharing in land registration using Hyperledger Fabric blockchain. Digital, 2(2), 6. https://doi.org/10.3390/digital2020006",
        "Ali, T., Nadeem, A., Alzahrani, A., & Jan, S. (2020). A transparent and trusted property registration system on permissioned blockchain. 2019 International Conference on Advances in the Emerging Computing Technologies. https://doi.org/10.1109/AECT47998.2020.9194222",
    ])
    path = OUT / "BLMS_Systematic_Literature_Review.docx"
    doc.save(path)
    return path


def build_results():
    doc = base_document("Implementation Results Report")
    title_page(
        doc,
        "Implementation Results Report",
        "Verified results of the locally runnable BLMS research prototype",
        "Lecturer Submission",
    )
    doc.add_heading("Executive summary", level=1)
    paragraph(doc, "The Blockchain-Based Land Management System (BLMS) has been implemented as a locally runnable academic prototype and deployed as a static browser application for demonstration. Its default path does not require Docker, a database service, IPFS, Hyperledger Fabric, an API server, environment variables, or external identity services. The deployed application uses browser-local storage and clearly synthetic development identities. This report presents only verified functionality and test evidence. It also records the boundary between the working local prototype and the remaining Hyperledger Fabric integration work.")
    add_callout(doc, "Implementation status.", "The local workflow is implemented and verified. The Go chaincode source is present as a Fabric research artifact. A live two-organisation Fabric network, Node Gateway adapter, live IPFS node, NIMC integration, Caliper benchmark, and expert evaluation are not yet implemented or executed.")

    doc.add_heading("1. Project objective and delivery approach", level=1)
    paragraph(doc, "BLMS addresses the need for more secure and transparent land-administration workflows by modelling land registration, ownership, document verification, sale listing, simulated transfer, government levy management, audit history, and revenue summary. The delivery approach was phased so the project could be tested locally and deployed for a supervisor demonstration without the cost and operational complexity of a blockchain network.")
    add_table(doc, ["Phase", "Delivered outcome", "Verification status"], [
        ["1. Foundation", "TypeScript monorepo, shared data contracts, API, identity gateway, and project documentation.", "Type-checked."],
        ["2. Local workflow", "Authentication, roles, parcels, registration, marketplace, transfer, audit, and revenue workflow.", "API tests passed."],
        ["3. Identity and documents", "Synthetic identity behaviour, document content hashing, and MATCH/MISMATCH verification.", "Identity and API tests passed."],
        ["4. Fabric research artifact", "Go land-chaincode source with land-life-cycle checks and role-based operations.", "Source added; network deployment is pending."],
        ["5. Offline/static delivery", "Browser-local adapter, responsive dashboard, Netlify-ready static build with no environment variables.", "Production web build passed."],
    ], [1440, 4920, 3000])

    doc.add_heading("2. Implemented architecture", level=1)
    paragraph(doc, "The implemented system uses a React/Vite frontend as the primary user interface. In its default offline mode, an in-browser adapter performs the same visible land-management workflow using localStorage. An optional Express API and mock ledger are retained for local API-mode testing. The project also contains a synthetic identity gateway and a Go chaincode package intended for a later Hyperledger Fabric deployment.")
    add_table(doc, ["Component", "Technology", "Role in the current implementation"], [
        ["Web dashboard", "React, TypeScript, Vite, React Router, React Query", "Responsive dashboard, role-aware navigation, forms, marketplace, history, and system status."],
        ["Offline data adapter", "Browser APIs and localStorage", "Default no-service workflow for authentication, parcels, transactions, audit events, and document hashes."],
        ["Optional API", "Node.js, Express, TypeScript, Zod, JWT, bcrypt", "Local API-mode validation, role checks, health endpoints, and mock-ledger access."],
        ["Identity gateway", "Node.js/Express", "Synthetic NIN verification with masked logging; no live government identity connection."],
        ["Document integrity", "SHA-256 browser digest and content identifier", "Detects whether a candidate PDF, PNG, or JPEG matches the recorded file content."],
        ["Fabric artifact", "Go chaincode / Hyperledger Fabric contract APIs", "Research-track contract for access control, duplicate checks, listing, transfer, levy, history, and events; not deployed."],
        ["Static deployment", "Netlify", "Hosts the browser-local frontend without environment variables or external backend services."],
    ], [1980, 2520, 4860])

    doc.add_heading("3. Functional results", level=1)
    doc.add_heading("3.1 Authentication and role-aware access", level=2)
    paragraph(doc, "The prototype provides synthetic development accounts for citizen, government-officer, and administrator roles. The interface adapts available actions by role, while the optional API tests validate access restrictions for protected actions. These accounts are expressly for local demonstration; they are not real users or government identities.")
    doc.add_heading("3.2 Land lifecycle", level=2)
    paragraph(doc, "A citizen can register a parcel with a parcel number, title number, location, size, and assessed value. Duplicate parcel or title values are rejected. The current owner can list or cancel a listing; another citizen can complete a simulated purchase. The transaction history and audit log record the workflow. The purchase result is clearly labelled simulated and moves no real funds.")
    doc.add_heading("3.3 Document verification", level=2)
    paragraph(doc, "The local workflow accepts PDF, PNG, and JPEG files up to 5 MB after basic content-signature checks. It calculates a SHA-256-based content identifier and stores document content locally. Re-uploading the identical file returns MATCH; altered content returns MISMATCH. The mechanism demonstrates content integrity; it does not constitute a production document vault or legal evidence service.")
    doc.add_heading("3.4 Government tools and auditability", level=2)
    paragraph(doc, "Government and administrator roles can update levies and access revenue summary data. The audit view records successful activity including actor, organisation, time, action, property reference, and detail. In offline mode these records reside in the browser and are therefore a demonstration of the workflow, not an independent tamper-resistant ledger.")
    doc.add_heading("3.5 Usability and deployment", level=2)
    paragraph(doc, "The dashboard was redesigned as a responsive, dark registry console. On smaller screens the sidebar becomes a slide-in menu with an overlay, close action, and Escape-key support. The application is available as a Netlify static deployment and can be run locally using Node.js and npm alone.")

    doc.add_heading("4. Verification evidence", level=1)
    paragraph(doc, "The following checks were run in the project workspace on 18 July 2026. The combined workspace test command was avoided in the final evidence table because empty Node test workspaces can remain open in this environment; the two test-bearing workspaces were executed directly.")
    add_table(doc, ["Check", "Observed result", "Meaning"], [
        ["npm run typecheck", "Passed for shared package, API, identity gateway, and web workspace.", "TypeScript interfaces and implementation compile without type errors."],
        ["API tests", "5 passed, 0 failed.", "Health mode, login/land query, purchase/fee simulation, government restriction, and document verification were exercised."],
        ["Identity-gateway tests", "3 passed, 0 failed.", "Synthetic-record health, masked synthetic verification, and status-key restriction were exercised."],
        ["Production web build", "Passed: Vite transformed 92 modules and produced the static build.", "The frontend can be packaged for static hosting."],
        ["Manual browser check", "Responsive sidebar opened, navigated, and closed as expected in the local browser check.", "The mobile navigation behaviour was manually verified."],
    ], [2340, 3510, 3510])

    doc.add_heading("5. Local run and deployment result", level=1)
    paragraph(doc, "For a local demonstration, install dependencies and run the offline command below. The application opens at http://localhost:5173 and stores its state in the browser only.")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("npm install\nnpm run dev:offline")
    set_run_font(r, size=10.5, color=BLUE)
    p.style = doc.styles["Normal"]
    paragraph(doc, "The static deployment is available at https://blms-local-registry.netlify.app. It works without environment variables because the build defaults to the browser-local adapter when no API URL is supplied. It is suitable for prototype review and user-interface testing, not for operational land administration.")

    doc.add_heading("6. Hyperledger Fabric contribution and boundary", level=1)
    paragraph(doc, "Hyperledger Fabric remains the principal blockchain requirement of the research project. The repository now contains Go chaincode that represents the intended ledger rules: certificate/attribute-based identity checks, duplicate parcel and title checks, owner-only listings, government-only levy updates, purchase rules, simulated fee allocation, document-CID comparison, history retrieval, and a LandTransferred event. This contributes a concrete blockchain-domain artifact rather than only a conceptual diagram.")
    paragraph(doc, "However, adding the chaincode does not mean that the web application is currently using a Fabric ledger. The default application deliberately stays independent so it can work without Docker or any external service. A real Fabric result requires a two-organisation network, certificates, channel policies, peer/orderer configuration, chaincode packaging and deployment, and a Gateway adapter that replaces the local/mock ledger behind the same application contract.")

    doc.add_heading("7. Limitations and next implementation phase", level=1)
    add_table(doc, ["Not yet implemented", "Why it matters", "Required next step"], [
        ["Live Fabric network and Gateway adapter", "Needed to demonstrate distributed endorsement, peer state, and real Fabric transaction submission.", "Deploy the two-organisation network locally, package Go chaincode, and implement FabricLedgerService."],
        ["Live IPFS/Kubo service", "Needed for networked document retrieval and pinning beyond browser storage.", "Run a local Kubo node and test upload/retrieval policy."],
        ["NIMC/NIN production integration", "Needed to verify real identities and meet privacy/legal obligations.", "Obtain formal authority, API access, data-protection review, and consent/retention design."],
        ["Caliper performance benchmark", "Needed to make any throughput/latency claim.", "Create workloads, record hardware/network conditions, run repeated trials, and publish raw results."],
        ["Expert or user evaluation", "Needed to claim usability, acceptance, or perceived trust outcomes.", "Obtain ethics/supervisor approval and run a documented study."],
    ], [2700, 3330, 3330])
    add_callout(doc, "Academic integrity note.", "The completed prototype demonstrates locally testable workflows and a static deployment. It does not demonstrate a production blockchain, real identity verification, legal title transfer, real payment settlement, or measured performance superiority.")

    doc.add_heading("8. Conclusion", level=1)
    paragraph(doc, "The BLMS implementation has produced a working, deployable local prototype that maps directly to the land-administration problems identified in the review: it improves demonstrability of registration rules, transaction traceability, document-content verification, role separation, and audit visibility. The architecture deliberately allows the application to function without Docker or external services, which makes it practical for academic demonstration and Netlify testing. The Go chaincode provides a credible starting point for the required Hyperledger Fabric research track. The next claimable research result should follow only after the Fabric network, integration tests, and transparent evaluation protocol have been completed.")
    add_references(doc, [
        "Project source code and verification run, BLMS repository, accessed 18 July 2026.",
        "BLMS README and phased implementation plan, local project documentation, accessed 18 July 2026.",
        "Netlify static deployment: https://blms-local-registry.netlify.app (prototype demonstration site).",
        "Androulaki, E., Barger, A., Bortnikov, V., et al. (2018). Hyperledger Fabric: A distributed operating system for permissioned blockchains. Proceedings of the Thirteenth EuroSys Conference. https://doi.org/10.1145/3190508.3190538",
    ])
    path = OUT / "BLMS_Implementation_Results_Report.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_slr())
    print(build_results())
