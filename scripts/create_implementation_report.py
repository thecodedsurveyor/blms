from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "BLMS_Implementation_and_Technology_Report.docx"

NAVY = "17324D"
BLUE = "2E74B5"
MUTED = "5E6A75"
PALE = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "1A1A1A"
USABLE_DXA = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for index, width in enumerate(widths):
        if index < len(grid.gridCol_lst):
            grid.gridCol_lst[index].set(qn("w:w"), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def border_bottom(paragraph, color=BLUE, size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_run(run, *, size=11, color=BLACK, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_para(doc, text="", *, style=None, size=None, color=None, bold=False, italic=False, align=None, before=None, after=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if text:
        run = p.add_run(text)
        set_run(run, size=size or 11, color=color or BLACK, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [USABLE_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{label}: ")
    set_run(run, size=10.5, color=NAVY, bold=True)
    run = p.add_run(text)
    set_run(run, size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PALE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run(r, size=9.5, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run(r, size=9.5, color=BLACK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(4)
    run = header.add_run("BLMS | IMPLEMENTATION AND TECHNOLOGY REPORT")
    set_run(run, size=8.5, color=MUTED, bold=True)
    border_bottom(header, color="D7DBE2", size="4", space="4")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Academic research prototype | July 2026")
    set_run(run, size=8.5, color=MUTED)


def build_report():
    doc = Document()
    configure_document(doc)

    # Editorial cover for a formal research report.
    add_para(doc, "PROJECT IMPLEMENTATION REPORT", size=11, color=BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_para(doc, "Blockchain-Based Land Management System (BLMS)", size=27, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "Implementation Status, Architecture, and Technology Rationale", size=15, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
    add_para(doc, "Prepared for project and research supervision", size=11, color=BLACK, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "Nigeria-focused academic prototype", size=11, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=50)
    add_para(doc, "Report date: 14 July 2026", size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "Classification: Academic prototype - synthetic data only", size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    add_heading(doc, "1. Executive Summary")
    add_para(doc, "The Blockchain-Based Land Management System (BLMS) is a research prototype designed to demonstrate how digital land registration, ownership transfer, document verification, and government oversight can be organised in a transparent and traceable workflow. The project is focused on the Nigerian land administration context, but it does not connect to NIMC, a government registry, banks, payment rails, or any production blockchain.")
    add_para(doc, "The main implementation is a browser-based, static application that can run locally or on Netlify without Docker, environment variables, a database, or external services. It provides role-aware workflows for citizens, government officers, and administrators. A separate Node.js API path, synthetic identity gateway, document vault, and Hyperledger Fabric Go chaincode are included as optional research and integration layers.")
    add_callout(doc, "Key conclusion", "The project demonstrates the full land-management user journey locally while keeping the production-like blockchain integration clearly separated and honestly marked as future integration work.")

    add_heading(doc, "2. Problem Addressed and Project Objective")
    add_para(doc, "Land administration often depends on fragmented records, manual verification, and processes that are difficult for citizens and government officers to audit. BLMS addresses this research problem by modelling a single workflow in which parcel information, ownership, document evidence, listings, transfers, levies, and audit events are visible through controlled roles.")
    add_para(doc, "The objective is not to claim a live national land registry. Instead, the objective is to produce a credible and demonstrable system architecture that shows where blockchain concepts, identity checks, content addressing, role-based control, and auditability fit into a land-management solution.")

    add_heading(doc, "3. What Has Been Implemented")
    add_table(doc, ["Area", "Delivered implementation", "Current status"], [
        ("Offline web application", "Role-aware dashboard, responsive navigation, parcel registration, search, marketplace, parcel history, levies, revenue, audit, profile, health, and reset flows.", "Implemented and deployed"),
        ("Browser-local data layer", "Authentication simulation, land lifecycle, transaction log, audit log, document storage metadata, and revenue simulation using browser localStorage.", "Implemented"),
        ("Document verification", "PDF, PNG, and JPEG validation; local SHA-256 content addressing; CID-style reference; MATCH/MISMATCH comparison.", "Implemented locally"),
        ("Synthetic identity", "Synthetic NIN records, masked NIN handling, local verification behaviour, and health/status endpoints.", "Implemented; no NIMC connection"),
        ("Optional Node.js API", "Express routes, validation, role checks, HTTP-only cookie sessions, mock-ledger adapter, audit logging, and file-upload handling.", "Implemented for local use"),
        ("Hyperledger Fabric chaincode", "Go contract for land rules, certificate attributes, ownership rules, transfers, levies, document CID references, and events.", "Source implemented; deployment pending"),
        ("Static deployment", "Vite production build, Netlify configuration, SPA route fallback, no required environment variables.", "Implemented and deployed"),
    ], [1800, 5200, 2360])
    add_para(doc, "The static deployment intentionally starts with an empty parcel register. Development accounts remain available only to demonstrate the workflows; no sample parcels, transactions, documents, or audit records are shipped as live-looking data.", italic=True, color=MUTED, size=10)

    add_heading(doc, "4. Implemented User Workflows")
    add_heading(doc, "4.1 Citizen workflow", level=2)
    add_bullet(doc, "Sign in with a synthetic development identity.")
    add_bullet(doc, "Create a parcel record using parcel, title, location, survey, and assessed-value details.")
    add_bullet(doc, "Attach a permitted document and obtain a local content reference.")
    add_bullet(doc, "Verify an original or changed file against the recorded reference.")
    add_bullet(doc, "List a parcel, inspect marketplace listings, and perform a simulated purchase between citizen roles.")
    add_heading(doc, "4.2 Government and administrator workflow", level=2)
    add_bullet(doc, "Government roles can update outstanding levies and inspect simulated revenue and audit data.")
    add_bullet(doc, "Administrators can review local service status and reset browser-local demonstration state.")
    add_bullet(doc, "The interface uses role-aware navigation so each role sees the tools relevant to its responsibility.")

    add_heading(doc, "5. Architecture in Plain Language")
    add_para(doc, "BLMS uses a layered architecture so that the interface can work independently of the optional enterprise integration components. The React application talks either to a browser-local adapter (the default) or to the optional Express API. Both paths use the same domain concepts: users, land assets, transactions, documents, levies, and audit events. This makes it possible to demonstrate the system immediately while preserving a route toward a ledger-backed deployment.")
    add_table(doc, ["Layer", "Responsibility", "Why it matters"], [
        ("Presentation layer", "React dashboard and responsive interface.", "Provides a clear experience for citizens and public officers."),
        ("Application layer", "Browser adapter or Express API; validation; role checks; workflow rules.", "Keeps business actions consistent across local and future ledger modes."),
        ("Evidence layer", "Content-addressed document reference and verification.", "Shows whether a candidate file matches the document reference on the land record."),
        ("Ledger research layer", "Mock ledger today; Go Fabric chaincode for a future network.", "Separates demonstrable behaviour from blockchain network deployment complexity."),
    ], [1650, 3700, 4010])

    add_heading(doc, "6. Technology Choices and Rationale")
    add_table(doc, ["Technology", "How it is used", "Why it was chosen"], [
        ("React and TypeScript", "User interface, components, role-aware pages, and typed domain models.", "React supports modular screens; TypeScript reduces mistakes by checking the shape of land, user, and transaction data during development."),
        ("Vite", "Development server and optimised static production build.", "It is fast, simple, and produces a static bundle suitable for free hosting such as Netlify."),
        ("React Router", "Client-side routes for dashboards, parcels, documents, levies, audit, and profile pages.", "It enables a multi-page application experience while still deploying as a static single-page application."),
        ("TanStack Query", "Loads and refreshes land, history, transaction, and health data.", "It keeps server or local-adapter data consistent after actions such as registration, listing, and transfer."),
        ("React Hook Form and Zod", "Form handling and input validation in the user interface and API contract.", "They make parcel registration and financial values easier to validate before a workflow is accepted."),
        ("Node.js and Express", "Optional local API and REST endpoints.", "Express is lightweight for demonstrating API design, security middleware, role checks, uploads, and testing without requiring a heavy enterprise platform."),
        ("bcrypt, JWT, HTTP-only cookies", "Optional API authentication path.", "Passwords are hashed; signed tokens and HTTP-only cookies model safer local session handling than storing raw passwords in the browser."),
        ("localStorage and Web Crypto", "Default offline/static data mode and file hashing.", "They allow a complete local demonstration and Netlify deployment without Docker, databases, or environment variables."),
        ("SHA-256 content references", "Produces a CID-style reference for a document's bytes.", "Changing the file changes its reference, making document comparison understandable in a local research demo."),
        ("Hyperledger Fabric and Go", "Future ledger contract layer for land rules and organisational permissions.", "Fabric suits permissioned multi-organisation records; Go is a common and efficient language for Fabric chaincode."),
        ("Netlify", "Static hosting for the React build.", "It serves the browser-only application at no cost for testing and does not require secrets or a running backend for the offline mode."),
        ("IBM Plex Sans and IBM Plex Mono", "Bundled interface and data typography.", "The fonts provide legible professional text and a distinct monospaced treatment for IDs, CIDs, and ledger-oriented data."),
    ], [1700, 3100, 4560])

    add_heading(doc, "7. Security and Integrity Measures")
    add_para(doc, "The prototype includes practical safeguards appropriate for an academic demonstration. These controls do not make the system production-ready, but they show the intended direction of a secure implementation.")
    add_bullet(doc, "Role checks distinguish citizens, government officers, and administrators.")
    add_bullet(doc, "The API mode hashes passwords with bcrypt and uses short-lived JWTs in HTTP-only cookies.")
    add_bullet(doc, "The API includes Helmet security headers, request-rate limiting, request IDs, input validation, and controlled upload limits.")
    add_bullet(doc, "Financial values are handled as integer kobo amounts in the application and chaincode design to avoid decimal-rounding errors.")
    add_bullet(doc, "Document uploads are restricted to validated PDF, PNG, and JPEG files, with a 5 MB maximum in the local vault.")
    add_bullet(doc, "The Fabric chaincode design keeps document bytes off-chain and records only safe metadata and a content reference.")

    add_heading(doc, "8. Hyperledger Fabric Position in the Project")
    add_para(doc, "Hyperledger Fabric is included because land administration is a permissioned environment: citizens, government bodies, and possibly other institutions need controlled access rather than an anonymous public blockchain. The Go chaincode defines the rules that should be enforced consistently when a Fabric network is later introduced.")
    add_table(doc, ["Chaincode rule", "Purpose"], [
        ("Certificate-attribute identity checks", "Ensures that state-changing actions use the caller's registered identity rather than an owner value supplied by a client."),
        ("Duplicate parcel and title checks", "Prevents a parcel or title identifier from being registered twice."),
        ("Owner-only listing", "Prevents a non-owner from offering a parcel for sale."),
        ("Atomic purchase logic", "Changes ownership, clears the listing, calculates a simulated 5% fee, and emits an event as one ledger operation."),
        ("Government-only levy updates", "Restricts levy changes to the designated government organisation."),
        ("History and events", "Supports later traceability of land changes and transfer activity."),
    ], [3150, 6210])
    add_callout(doc, "Important status", "The Go chaincode source is present, but a Fabric network, certificate authority setup, Node Gateway adapter, and actual chaincode deployment have not yet been completed or verified in this workspace. The application therefore does not claim to be running on Fabric today.")

    add_heading(doc, "9. Testing and Verification Performed")
    add_para(doc, "The following checks have been run during implementation:")
    add_bullet(doc, "Workspace TypeScript type checking completed successfully.")
    add_bullet(doc, "The Vite production build for the web application completed successfully before Netlify deployment.")
    add_bullet(doc, "API test suite completed with five passing tests, covering health mode, authentication/query flow, simulated purchase, government revenue access, and document verification.")
    add_bullet(doc, "Synthetic identity gateway test suite completed with three passing tests, covering health, masked verification, and protected status access.")
    add_bullet(doc, "Browser checks verified the mobile navigation drawer opens and closes after navigation, and the static deployment redirects client-side routes correctly.")
    add_para(doc, "Fabric deployment tests have not been claimed because Go and a Fabric/Docker network were not installed or run in this workspace.", italic=True, color=MUTED, size=10)

    add_heading(doc, "10. Known Limitations and Research Boundaries")
    add_bullet(doc, "All users, NINs, land values, payments, organisations, and documents are synthetic academic data.")
    add_bullet(doc, "The default static mode stores state only in the browser. It is not shared between users or devices and is not an immutable blockchain ledger.")
    add_bullet(doc, "No live NIMC verification, government registry connection, banking integration, payment processing, or legal conveyancing process is included.")
    add_bullet(doc, "The transfer and 5% revenue allocation are simulations; no money moves.")
    add_bullet(doc, "The Fabric chaincode is a research artefact awaiting network deployment and Gateway integration.")
    add_bullet(doc, "This work is suitable for demonstration and research discussion, not production land administration.")

    add_heading(doc, "11. Recommended Next Steps")
    add_bullet(doc, "Create the optional two-organisation Fabric network and deploy the Go chaincode with the intended endorsement policy.")
    add_bullet(doc, "Implement the Fabric Gateway adapter behind the existing LedgerService interface, then run the same API contract tests in mock and Fabric modes.")
    add_bullet(doc, "Replace the synthetic identity gateway only after obtaining lawful approval and an authorised integration path to a real identity provider.")
    add_bullet(doc, "Introduce a managed database and secure off-chain document storage for multi-user persistence before any real deployment.")
    add_bullet(doc, "Add end-to-end browser automation, performance testing, and a formal threat model before a larger pilot.")
    add_bullet(doc, "Conduct usability evaluation with citizens, surveyors, and land-administration officers to validate workflow assumptions.")

    add_heading(doc, "12. Conclusion")
    add_para(doc, "BLMS has achieved its immediate research objective: it provides a complete, locally demonstrable land-management workflow and a credible technical path toward a permissioned blockchain implementation. The static deployment makes the prototype easy to test, while the optional API and chaincode layers document how the system can evolve beyond a browser-only demonstration. The most important point for evaluation is that the project separates implemented functionality from planned infrastructure and does not overstate its current level of production readiness.")

    add_heading(doc, "Appendix A: Source Basis for This Report")
    add_para(doc, "This report was prepared from the current BLMS source tree and implementation documentation, including the project README, offline walkthrough, phased implementation plan, web/API package manifests, Express routes, local document vault, synthetic identity gateway, and the Go land-chaincode README and source. Verification statements reflect checks actually run in the local workspace during implementation.", size=10, color=MUTED)

    doc.core_properties.title = "BLMS Implementation and Technology Report"
    doc.core_properties.subject = "Academic research prototype implementation report"
    doc.core_properties.author = "BLMS Project Team"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
